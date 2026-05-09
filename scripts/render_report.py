"""Render Markdown reports to HTML and optionally DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path


def render_html(input_path: Path, html_output: Path) -> None:
    import markdown

    md = input_path.read_text(encoding="utf-8")
    body = markdown.markdown(md, extensions=["tables", "fenced_code"])
    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{input_path.stem}</title>
  <style>
    body {{ font-family: "Microsoft YaHei", Arial, sans-serif; line-height: 1.7; max-width: 960px; margin: 32px auto; padding: 0 20px; color: #222; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid #d0d7de; padding: 8px 10px; vertical-align: top; }}
    th {{ background: #f6f8fa; }}
    blockquote {{ margin: 12px 0; padding: 8px 14px; border-left: 4px solid #8c959f; background: #f6f8fa; }}
    code {{ background: #f6f8fa; padding: 2px 4px; border-radius: 4px; }}
  </style>
</head>
<body>
{body}
</body>
</html>
"""
    html_output.parent.mkdir(parents=True, exist_ok=True)
    html_output.write_text(html, encoding="utf-8")


def render_docx(input_path: Path, docx_output: Path) -> None:
    from docx import Document

    doc = Document()
    for line in input_path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    docx_output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Render Markdown report.")
    parser.add_argument("--input", required=True)
    parser.add_argument("--html-output")
    parser.add_argument("--docx-output")
    args = parser.parse_args()
    input_path = Path(args.input)
    if args.html_output:
        render_html(input_path, Path(args.html_output))
    if args.docx_output:
        render_docx(input_path, Path(args.docx_output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

