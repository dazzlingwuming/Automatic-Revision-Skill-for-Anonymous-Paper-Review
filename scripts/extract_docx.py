"""Extract paragraph and table text from a DOCX file."""

from __future__ import annotations

import argparse
import json
from pathlib import Path


def extract_docx(input_path: Path, output_path: Path) -> dict:
    from docx import Document

    doc = Document(input_path)
    paragraphs = [
        {
            "index": index,
            "style": paragraph.style.name if paragraph.style else None,
            "text": paragraph.text,
        }
        for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.strip()
    ]
    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    data = {
        "source_file": str(input_path),
        "file_type": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
        "warnings": ["DOCX page numbers are not reliably available."],
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return data


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract text from DOCX.")
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    extract_docx(Path(args.input), Path(args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

