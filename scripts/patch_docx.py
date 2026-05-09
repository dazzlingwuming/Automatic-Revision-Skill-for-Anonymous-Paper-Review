"""Create a DOCX revision suggestion draft."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))


def build_patch_docx(revision_plans_dir: Path, output_path: Path) -> None:
    from docx import Document

    plans = []
    for path in sorted(revision_plans_dir.glob("*.json")):
        plans.append(json.loads(path.read_text(encoding="utf-8")))

    doc = Document()
    doc.add_heading("论文修改建议版", level=1)
    doc.add_paragraph("本文件仅汇总拟修改建议，不表示已完成正文修改。")
    table = doc.add_table(rows=1, cols=6)
    headers = ["意见编号", "处理状态", "位置", "修改类型", "修改前摘录", "建议修改文本"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header

    for plan in plans:
        actions = plan.get("specific_actions", []) or [{}]
        for action in actions:
            location = action.get("location", {})
            row = table.add_row().cells
            row[0].text = plan.get("comment_id", "")
            row[1].text = plan.get("revision_status", "")
            row[2].text = f"{location.get('section') or ''} {location.get('page_range') or ''} {location.get('chunk_id') or ''}".strip()
            row[3].text = action.get("type", "")
            row[4].text = action.get("before_excerpt", "")
            row[5].text = action.get("after_proposed_text", "")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Create DOCX revision suggestion draft.")
    parser.add_argument("--input-docx", help="Original DOCX. When provided, suggestions are inserted into a copy.")
    parser.add_argument("--revision-plans-dir", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    if args.input_docx:
        from src.patching.docx_writer import patch_docx_with_revision_plans

        patch_docx_with_revision_plans(Path(args.input_docx), Path(args.revision_plans_dir), Path(args.output))
    else:
        build_patch_docx(Path(args.revision_plans_dir), Path(args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
