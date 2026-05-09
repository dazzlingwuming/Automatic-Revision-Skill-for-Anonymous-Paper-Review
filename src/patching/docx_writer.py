"""Create DOCX suggestion copies from revision plans."""

from __future__ import annotations

import json
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph


def load_revision_plans(revision_plans_dir: Path) -> list[dict]:
    return [json.loads(path.read_text(encoding="utf-8")) for path in sorted(revision_plans_dir.glob("*.json"))]


def patch_docx_with_revision_plans(input_docx: Path, revision_plans_dir: Path, output_path: Path) -> None:
    document = Document(input_docx)
    for plan in load_revision_plans(revision_plans_dir):
        for action in _plan_actions(plan):
            new_text = action.get("new_text") or action.get("after_proposed_text") or ""
            if not new_text.strip():
                continue
            anchor = action.get("anchor_text") or action.get("original_text") or action.get("before_excerpt") or ""
            paragraph = _find_anchor_paragraph(document, anchor)
            inserted_text = f"【建议新增 {plan.get('comment_id', '')}/{action.get('action_id', '')}】{new_text}"
            note_text = f"【批注说明】回应 {plan.get('comment_id', '')}：{action.get('rationale') or plan.get('overall_strategy') or plan.get('revision_strategy') or '请作者核对该处建议。'}"
            if paragraph is None:
                paragraph = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
                inserted_text = "【需人工确认位置】" + inserted_text
            inserted = _insert_paragraph_after(paragraph, inserted_text)
            _highlight_paragraph(inserted)
            _insert_paragraph_after(inserted, note_text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def apply_revision_plans_to_docx(input_docx: Path, revision_plans_dir: Path, output_path: Path) -> None:
    document = Document(input_docx)
    for plan in load_revision_plans(revision_plans_dir):
        for action in _plan_actions(plan):
            if action.get("requires_author_input"):
                continue
            action_type = action.get("type", "")
            new_text = action.get("new_text") or action.get("after_proposed_text") or ""
            if not new_text.strip():
                continue
            anchor = action.get("anchor_text") or action.get("original_text") or action.get("before_excerpt") or ""
            paragraph = _find_anchor_paragraph(document, anchor)
            if paragraph is None:
                continue
            if action_type in {"replace_paragraph", "rewrite_sentence"}:
                paragraph.clear()
                paragraph.add_run(new_text)
            elif action_type in {"insert_after_paragraph", "add", "rewrite"}:
                _insert_paragraph_after(paragraph, new_text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _plan_actions(plan: dict) -> list[dict]:
    if plan.get("actions"):
        return plan["actions"]
    return plan.get("specific_actions", [])


def _find_anchor_paragraph(document, anchor_text: str) -> Paragraph | None:
    clean_anchor = " ".join(anchor_text.split())
    if not clean_anchor:
        return None
    for paragraph in document.paragraphs:
        if clean_anchor in " ".join(paragraph.text.split()):
            return paragraph
    for paragraph in document.paragraphs:
        if clean_anchor[:30] and clean_anchor[:30] in " ".join(paragraph.text.split()):
            return paragraph
    return None


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.clear()
    inserted.add_run(text)
    return inserted


def _highlight_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
