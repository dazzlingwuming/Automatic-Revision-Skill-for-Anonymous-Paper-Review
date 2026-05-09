"""Word 文档解析模块"""

from pathlib import Path


class DocxParser:
    """Word 文档解析器"""

    def __init__(self, file_path: str | Path):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"Word 文件不存在: {file_path}")

    def extract_text(self) -> str:
        """提取全部文本"""
        from docx import Document
        doc = Document(self.file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    def get_paragraphs(self) -> list[dict]:
        """获取段落列表（含样式信息）"""
        from docx import Document
        doc = Document(self.file_path)
        return [
            {"text": p.text, "style": p.style.name if p.style else None}
            for p in doc.paragraphs if p.text.strip()
        ]

    def get_toc(self) -> list[str]:
        """提取目录（通过标题样式识别）"""
        from docx import Document
        doc = Document(self.file_path)
        toc = []
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith("Heading"):
                level = p.style.name.replace("Heading ", "")
                indent = "  " * (int(level) - 1) if level.isdigit() else ""
                toc.append(f"{indent}{p.text}")
        return toc
