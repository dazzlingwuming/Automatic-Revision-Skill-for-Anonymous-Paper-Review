from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]


def run_script(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run([sys.executable, *args], cwd=ROOT, text=True, capture_output=True, check=False)


def test_patch_docx_inserts_highlighted_suggestions(tmp_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    input_docx = tmp_path / "paper.docx"
    plans_dir = tmp_path / "plans"
    output_docx = tmp_path / "outputs" / "修改建议版.docx"
    plans_dir.mkdir()

    doc = Document()
    doc.add_heading("第一章 绪论", level=1)
    doc.add_paragraph("本文介绍研究背景。")
    doc.save(input_docx)

    plan = {
        "comment_id": "R1-C001",
        "revision_status": "text_ready",
        "overall_strategy": "补充研究背景。",
        "actions": [
            {
                "action_id": "A1",
                "type": "insert_after_paragraph",
                "target": {"section_id": "sec_1", "section_title": "第一章 绪论"},
                "anchor_text": "本文介绍研究背景。",
                "original_text": "本文介绍研究背景。",
                "new_text": "为回应盲审专家关于研究背景展开不足的意见，本文进一步补充研究问题的现实来源、理论依据和论文后续章节之间的衔接关系。",
                "rationale": "回应背景不足。",
                "requires_author_input": False,
            }
        ],
        "reviewer_response": "感谢专家意见，本文拟补充研究背景说明。",
        "author_input_needed": [],
        "risks": [],
        "confidence": 0.9,
    }
    (plans_dir / "R1-C001.json").write_text(json.dumps(plan, ensure_ascii=False), encoding="utf-8")

    result = run_script(
        "scripts/patch_docx.py",
        "--input-docx",
        str(input_docx),
        "--revision-plans-dir",
        str(plans_dir),
        "--output",
        str(output_docx),
    )
    assert result.returncode == 0, result.stderr + result.stdout

    patched = Document(output_docx)
    texts = [paragraph.text for paragraph in patched.paragraphs]
    inserted = next(text for text in texts if "为回应盲审专家" in text)
    assert inserted.startswith("【建议新增 R1-C001/A1】")
    inserted_para = next(paragraph for paragraph in patched.paragraphs if "为回应盲审专家" in paragraph.text)
    assert inserted_para.runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW
