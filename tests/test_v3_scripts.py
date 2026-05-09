"""Tests for v3 deterministic scripts and schemas."""

from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]


def run_script(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, *args],
        cwd=ROOT,
        text=True,
        capture_output=True,
        check=False,
    )


def test_validate_review_comments_schema() -> None:
    result = run_script(
        "scripts/validate_json.py",
        "--schema",
        ".claude/skills/thesis-review-revision/schemas/review_comments.schema.json",
        "--input",
        "tests/fixtures/sample_review_comments.json",
    )
    assert result.returncode == 0, result.stderr + result.stdout


def test_extract_txt(tmp_path: Path) -> None:
    output = tmp_path / "review_raw.txt"
    result = run_script(
        "scripts/extract_txt.py",
        "--input",
        "tests/fixtures/sample_review.txt",
        "--output",
        str(output),
    )
    assert result.returncode == 0, result.stderr + result.stdout
    assert "论文对研究创新点" in output.read_text(encoding="utf-8")


def test_extract_txt_json_output(tmp_path: Path) -> None:
    output = tmp_path / "paper_raw.json"
    result = run_script(
        "scripts/extract_txt.py",
        "--input",
        "tests/fixtures/sample_paper.txt",
        "--output",
        str(output),
    )
    assert result.returncode == 0, result.stderr + result.stdout
    data = json.loads(output.read_text(encoding="utf-8"))
    assert data["file_type"] == "txt"
    assert "摘要" in data["text"]


def test_extract_docx(tmp_path: Path) -> None:
    from docx import Document

    input_path = tmp_path / "paper.docx"
    output = tmp_path / "paper_raw.json"
    doc = Document()
    doc.add_heading("第一章 绪论", level=1)
    doc.add_paragraph("本文介绍研究背景。")
    doc.save(input_path)
    result = run_script(
        "scripts/extract_docx.py",
        "--input",
        str(input_path),
        "--output",
        str(output),
    )
    assert result.returncode == 0, result.stderr + result.stdout
    data = json.loads(output.read_text(encoding="utf-8"))
    assert data["file_type"] == "docx"
    assert data["paragraphs"]


def test_extract_pdf(tmp_path: Path) -> None:
    import fitz

    input_path = tmp_path / "paper.pdf"
    output = tmp_path / "paper_raw.json"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "第一章 绪论")
    doc.save(input_path)
    doc.close()
    result = run_script(
        "scripts/extract_pdf.py",
        "--input",
        str(input_path),
        "--output",
        str(output),
    )
    assert result.returncode == 0, result.stderr + result.stdout
    data = json.loads(output.read_text(encoding="utf-8"))
    assert data["file_type"] == "pdf"
    assert data["pages"][0]["page_number"] == 1


def test_validate_all_core_examples() -> None:
    pairs = [
        ("paper_index.schema.json", "paper_index.example.json"),
        ("comment_mappings.schema.json", "comment_mappings.example.json"),
        ("revision_plan.schema.json", "revision_plan.example.json"),
    ]
    for schema, example in pairs:
        result = run_script(
            "scripts/validate_json.py",
            "--schema",
            f".claude/skills/thesis-review-revision/schemas/{schema}",
            "--input",
            f".claude/skills/thesis-review-revision/examples/{example}",
        )
        assert result.returncode == 0, result.stderr + result.stdout


def test_chunk_paper_outputs_metadata(tmp_path: Path) -> None:
    chunks_dir = tmp_path / "chunks"
    metadata = tmp_path / "artifacts" / "paper_chunks.json"
    result = run_script(
        "scripts/chunk_paper.py",
        "--input",
        "tests/fixtures/sample_paper_raw.json",
        "--chunks-dir",
        str(chunks_dir),
        "--metadata-output",
        str(metadata),
        "--target-chars",
        "80",
    )
    assert result.returncode == 0, result.stderr + result.stdout
    data = json.loads(metadata.read_text(encoding="utf-8"))
    assert data["chunks"]
    assert (chunks_dir / "ch_0001.txt").exists()


def test_format_checker_outputs_issues(tmp_path: Path) -> None:
    output = tmp_path / "format_issues.json"
    result = run_script(
        "scripts/format_checker.py",
        "--paper-index",
        "tests/fixtures/sample_paper_index.json",
        "--output",
        str(output),
    )
    assert result.returncode == 0, result.stderr + result.stdout
    data = json.loads(output.read_text(encoding="utf-8"))
    assert "issues" in data


def test_render_report_html(tmp_path: Path) -> None:
    html = tmp_path / "report.html"
    result = run_script(
        "scripts/render_report.py",
        "--input",
        "tests/fixtures/sample_report.md",
        "--html-output",
        str(html),
    )
    assert result.returncode == 0, result.stderr + result.stdout
    assert "<html" in html.read_text(encoding="utf-8")


def test_build_report_includes_before_after_examples(tmp_path: Path) -> None:
    plans_dir = tmp_path / "plans"
    plans_dir.mkdir()
    plan = json.loads((ROOT / "tests/fixtures/sample_revision_plan.json").read_text(encoding="utf-8"))
    (plans_dir / "R1-C001.json").write_text(json.dumps(plan, ensure_ascii=False), encoding="utf-8")
    output_dir = tmp_path / "outputs"
    result = run_script(
        "scripts/build_report.py",
        "--review-comments",
        "tests/fixtures/sample_review_comments.json",
        "--revision-plans-dir",
        str(plans_dir),
        "--output-dir",
        str(output_dir),
        "--paper-file",
        "paper.pdf",
        "--review-file",
        "review.pdf",
    )
    assert result.returncode == 0, result.stderr + result.stdout
    report = (output_dir / "修改报告.md").read_text(encoding="utf-8")
    assert "##### 修改前" in report
    assert "本文贡献需要进一步凝练" in report
