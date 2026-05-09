from __future__ import annotations

import json
import subprocess
import sys
import base64
from pathlib import Path

from jsonschema import Draft202012Validator


ROOT = Path(__file__).resolve().parents[1]


def run_script(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run([sys.executable, *args], cwd=ROOT, text=True, capture_output=True, check=False)


def make_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("第一章 绪论", level=1)
    doc.add_paragraph("本文介绍研究背景。")
    doc.add_paragraph("表1.1 实验参数设置")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "参数"
    table.rows[0].cells[1].text = "取值"
    table.rows[1].cells[0].text = "alpha"
    table.rows[1].cells[1].text = "0.1"
    doc.add_heading("1.1 研究内容", level=2)
    doc.add_paragraph("本文提出面向盲审修改的 DOCX-first 工作流。")
    doc.save(path)


def make_docx_with_picture(path: Path, image_path: Path) -> None:
    from docx import Document

    image_path.write_bytes(base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="))
    doc = Document()
    doc.add_heading("第一章 绪论", level=1)
    doc.add_paragraph("图1.1 技术路线图")
    doc.add_picture(str(image_path))
    doc.add_paragraph("本文围绕技术路线展开说明。")
    doc.save(path)


def test_docx_blocks_preserve_body_order_and_locators(tmp_path: Path) -> None:
    from src.ingestion.docx_blocks import extract_docx_blocks

    input_path = tmp_path / "paper.docx"
    make_docx(input_path)

    data = extract_docx_blocks(input_path)
    blocks = data["blocks"]

    assert [block["type"] for block in blocks[:4]] == ["heading", "paragraph", "table_caption", "table"]
    assert blocks[0]["style_name"].startswith("Heading") or blocks[0]["style_name"].startswith("标题")
    assert blocks[0]["docx_locator"]["body_index"] == 1
    assert blocks[3]["docx_locator"]["xml_tag"] == "w:tbl"
    assert blocks[3]["table"]["rows"] == [["参数", "取值"], ["alpha", "0.1"]]


def test_ingest_docx_builds_sections_and_markdown(tmp_path: Path) -> None:
    from src.docx_ingestion.ingest import ingest_docx

    input_path = tmp_path / "paper.docx"
    out = tmp_path / "run"
    make_docx(input_path)

    ingest_docx(input_path, out, title="测试论文")

    structure = json.loads((out / "paper" / "paper_structure.json").read_text(encoding="utf-8"))
    paper_md = (out / "paper" / "paper.md").read_text(encoding="utf-8")

    assert [section["section_id"] for section in structure["sections"][:2]] == ["sec_1", "sec_1_1"]
    assert (out / "paper" / "sections" / "sec_1.md").exists()
    assert "# 第一章 绪论" in paper_md
    assert "表1.1 实验参数设置" in paper_md
    assert "| 参数 | 取值 |" in paper_md


def test_docx_heading_styles_without_visible_numbers_get_hierarchical_ids(tmp_path: Path) -> None:
    from docx import Document
    from src.docx_ingestion.ingest import ingest_docx

    input_path = tmp_path / "paper.docx"
    out = tmp_path / "run"
    doc = Document()
    doc.add_paragraph("封面内容")
    doc.add_heading("绪论", level=1)
    doc.add_heading("研究背景", level=2)
    doc.add_paragraph("研究背景正文。")
    doc.add_heading("方法", level=1)
    doc.add_heading("模型构建", level=2)
    doc.save(input_path)

    ingest_docx(input_path, out)

    structure = json.loads((out / "paper" / "paper_structure.json").read_text(encoding="utf-8"))
    ids = [section["section_id"] for section in structure["sections"]]
    assert ids == ["sec_0001", "sec_1", "sec_1_1", "sec_2", "sec_2_1"]


def test_ingest_docx_extracts_table_assets(tmp_path: Path) -> None:
    from src.docx_ingestion.ingest import ingest_docx

    input_path = tmp_path / "paper.docx"
    out = tmp_path / "run"
    make_docx(input_path)

    ingest_docx(input_path, out)

    catalog = json.loads((out / "assets" / "asset_catalog.json").read_text(encoding="utf-8"))
    tables = [asset for asset in catalog["assets"] if asset["asset_type"] == "table"]

    assert len(tables) == 1
    assert tables[0]["asset_id"] == "tab_1_1"
    assert tables[0]["label"] == "表1.1"
    assert Path(tables[0]["markdown_path"]).exists()
    assert Path(tables[0]["csv_path"]).exists()
    assert tables[0]["quality"]["has_caption"] is True


def test_ingest_docx_extracts_figure_assets(tmp_path: Path) -> None:
    from src.docx_ingestion.ingest import ingest_docx

    input_path = tmp_path / "paper.docx"
    source_image = tmp_path / "source.png"
    out = tmp_path / "run"
    make_docx_with_picture(input_path, source_image)

    ingest_docx(input_path, out)

    catalog = json.loads((out / "assets" / "asset_catalog.json").read_text(encoding="utf-8"))
    figures = [asset for asset in catalog["assets"] if asset["asset_type"] == "figure"]

    assert len(figures) == 1
    assert figures[0]["asset_id"] == "fig_1_1"
    assert figures[0]["label"] == "图1.1"
    assert Path(figures[0]["image_path"]).exists()
    assert figures[0]["caption"] == "图1.1 技术路线图"


def test_ingest_docx_script_outputs_schema_valid_artifacts(tmp_path: Path) -> None:
    input_path = tmp_path / "paper.docx"
    out = tmp_path / "run"
    make_docx(input_path)

    result = run_script("scripts/ingest_docx.py", "--paper", str(input_path), "--out", str(out), "--title", "测试论文")
    assert result.returncode == 0, result.stderr + result.stdout

    schema_pairs = [
        ("paper_blocks.schema.json", out / "paper" / "paper_blocks.json"),
        ("paper_structure.schema.json", out / "paper" / "paper_structure.json"),
        ("asset_catalog.schema.json", out / "assets" / "asset_catalog.json"),
    ]
    for schema_name, artifact_path in schema_pairs:
        schema = json.loads((ROOT / ".claude" / "skills" / "thesis-review-revision" / "schemas" / schema_name).read_text(encoding="utf-8"))
        artifact = json.loads(artifact_path.read_text(encoding="utf-8"))
        Draft202012Validator(schema).validate(artifact)
